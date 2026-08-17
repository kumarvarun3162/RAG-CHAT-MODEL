import os
from docx import Document


def load_docx(filepath):
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    return Document(filepath)


def extract_paragraphs(doc):
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) < 4:          
            continue
        lines.append(text)
    return lines


def extract_tables(doc):
    rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
    return rows


def clean_lines(lines):
    cleaned = []
    for line in lines:
        if line.startswith("RP-") and "|" in line:
            continue
        if line.startswith("3GPP TS") or line.startswith("3GPP TR"):
            continue
        if line.lower().startswith("release") and len(line) < 25:
            continue
        if line.isdigit():
            continue
        cleaned.append(line)
    return cleaned


def ingest_docx(filepath):
    print(f"Reading: {filepath}")
    doc = load_docx(filepath)

    para_lines  = extract_paragraphs(doc)
    table_lines = extract_tables(doc)

    print(f"  Paragraphs: {len(para_lines)}")
    print(f"  Table rows: {len(table_lines)}")

    all_lines = para_lines + table_lines
    all_lines = clean_lines(all_lines)

    full_text = "\n".join(all_lines)
    print(f"  Total characters after cleaning: {len(full_text)}")
    return full_text


def save_text(text, output_path):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"Saved to: {output_path}")


if __name__ == "__main__":
    input_file  = "data/raw/38300-j30.docx"
    output_file = "data/extracted_text.txt"

    text = ingest_docx(input_file)
    save_text(text, output_file)

    print("\n--- Preview (first 600 chars) ---")
    print(text[:600])